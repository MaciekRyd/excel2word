import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import base64

st.set_page_config(page_title="Excel ➜ Word Generator", layout="wide")

st.title("📄 Generator dokumentów Word na podstawie Excela")

# 1. Wczytaj plik Excel
excel_file = st.file_uploader("Wczytaj plik Excel (.xlsx)", type=["xlsx"])

if excel_file:
    try:
        df = pd.read_excel(excel_file)
        st.success("Plik wczytany poprawnie.")
    except Exception as e:
        st.error(f"Błąd wczytywania: {e}")
        st.stop()

    st.subheader("Dane z Excela")
    selected_indices = st.multiselect("Wybierz wiersze do eksportu:", df.index.tolist(), format_func=lambda x: str(df.iloc[x].to_dict()))
    selected_rows = df.iloc[selected_indices]

    # 2. Wczytaj szablon Word
    template_file = st.file_uploader("Wczytaj szablon Word (.docx)", type=["docx"])

    if template_file and not selected_rows.empty:
        # 3. Generuj dokument Word
        if st.button("📤 Generuj dokument Word"):
            try:
                base_template = Document(template_file)
                result_doc = Document()

                for idx, row in selected_rows.iterrows():
                    temp_doc = Document(template_file)  # Zaczynaj od szablonu dla każdego rekordu
                    replacements = row.to_dict()

                    for p in temp_doc.paragraphs:
                        for key, value in replacements.items():
                            if f"{{{{{key}}}}}" in p.text:
                                inline = p.runs
                                for i in inline:
                                    if f"{{{{{key}}}}}" in i.text:
                                        i.text = i.text.replace(f"{{{{{key}}}}}", str(value))

                    # Dodaj zawartość tego szablonu do głównego dokumentu
                    for element in temp_doc.element.body:
                        result_doc.element.body.append(element)

                    result_doc.add_page_break()

                # 4. Zapisz dokument do bufora
                output = BytesIO()
                result_doc.save(output)
                output.seek(0)

                b64 = base64.b64encode(output.read()).decode()
                href = f'<a href="data:application/octet-stream;base64,{b64}" download="wynik.docx">📥 Pobierz dokument Word</a>'
                st.markdown(href, unsafe_allow_html=True)

            except Exception as e:
                st.error(f"Wystąpił błąd przy generowaniu dokumentu: {e}")
    elif template_file and selected_rows.empty:
        st.warning("Wybierz przynajmniej jeden rekord.")
